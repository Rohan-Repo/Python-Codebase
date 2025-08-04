# We need docx Python Module to work with DOCX Files so install it using pip
# pip install python-docx


from docx import Document

from pathlib import Path

# This is just a Dummy Value since Null-Check was not working
finalDoc = 'Inital-Doc'

# Find all .docx files in the specified path 
for fileInPath in Path('docx-files\\').glob('*.docx'):
    # Make Sure DOCX File is present
    if fileInPath:
        fileName = fileInPath.name

        print( '\n \t For Path : ', fileInPath, ' \t Name : ', fileName )

        # Open Each Document from the folder so have provided the filePath
        doc = Document( fileInPath )

        # Read all paragraphs from each document
        all_paras = doc.paragraphs
        print( '\t Len: ', len(all_paras) )
        
        # If the final document is not created, create it
        if finalDoc == 'Inital-Doc':
            finalDoc = Document()
        else:
        # If it is created open and append to it
            finalDoc = Document('finalDoc.docx')

        # Add Heading tp the document
        finalDoc.add_heading( 'DOCX Appended File!', level=1 )

        # Print Each Filename at the Start
        para = finalDoc.add_paragraph( '\n fileName:\n \t ' )
        # Each Filename in Bold
        para.add_run( fileName ).bold = True
        para.add_run( '\n' )

        # Add all paragraphs to the final doc
        for para in all_paras:
            finalDoc.add_paragraph( para.text )             
            # print( '\t', para.text )

        # print( '\n ---x-x-x-x-x-x-x-x-x-x-x--x \n' )
        # Add a Page Break after each document
        finalDoc.add_page_break()
        
        # Save the contents to the file
        finalDoc.save( 'finalDoc.docx' )

    else:
        print( 'DOCX File Not Found' )